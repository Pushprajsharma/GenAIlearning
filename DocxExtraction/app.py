from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
import requests
import google.generativeai as genai
import httpx
from groq import Groq
import docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


app = Flask(__name__)
CORS(app)

def extract_performance_requirements(file,actualRequirement):
    doc = Document(file)
    content = []
    for para in doc.paragraphs:
        content.append(para.text)
    
    detailedRequirment = "" 
    for text in content:
        detailedRequirment += text
    
    prompt = f"""
        You are a QA Automation Engineer specialized in writing test cases for the User Specific
        Requirements.

        I have a User Specific Requirements Specification (URS) document that outlines specific requirements
        for a software system. I want to generate detailed test cases for each requirement to ensure thorough validation.

        Understand the context of the application using {detailedRequirment}.

        Create well defined test cases for each requirement heading present in {actualRequirement} by using the context given above

        The test cases should be split and segregated for well defined result , Create Positive and Negative Test cases 

        The first Table should contain 4 rows refer below example:
        Test ID                |UAT_TS_001
        Test Description       |Test for record creation, assign numbers
        Requirement Reference  |U_MAN_0010, U_MAN_0020, U_MAN_0030, U_MAN_0040, U_MAN_0050, U_MAN_0060
        Pre-Requisites         |Availability of Test Data, Test Environment or tools etc.
                               |Tester must read script before execution,

        and then the steps in following format:
        Step    Step Description    Expected Results    Actual Results    Pass / Fail    Defect ID
        Ensure that the test cases are comprehensive covering edge cases, boundary values and failure scenarios.
        Provide additional details where necessary to make execution clear for testers.

        Segregate the test cases based on the requirement and requirement ID change the heading to Validation and not requirments
        specified in the detailedRequirement.

        Create sections for each requirement individually.

        Provide me the response in HTML format so I can directly render a table at UI
    """
    
    return getGeminiResponse(prompt)
    # return content


def extract_requirements_and_tables(doc_path):
    # Load the document
    doc = docx.Document(doc_path)
    requirements = {}
    current_heading = None

    # Iterate through each element in the document's body
    for element in doc.element.body:
        if isinstance(element, CT_P):  # Paragraph
            para = docx.text.paragraph.Paragraph(element, doc)

            # Check if the paragraph is a heading containing "Requirements"
            if para.style.name.startswith('Heading') and "Requirements" in para.text:
                current_heading = para.text.strip()
                requirements[current_heading] = {
                    'text': [],
                    'tables': []
                }
            elif current_heading:
                # Add the paragraph to the current heading's text list
                text = para.text.strip()
                if text:
                    requirements[current_heading]['text'].append(text)
        elif isinstance(element, CT_Tbl):  # Table
            table = docx.table.Table(element, doc)
            if current_heading:
                # Add the table to the current heading's tables list
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                requirements[current_heading]['tables'].append(table_data)

    # Remove headings without any text or tables
    requirements = {heading: content for heading, content in requirements.items() if content['text'] or content['tables']}
    
    output_content = ""
    for heading, content in requirements.items():
        output_content += f"{heading}:\n"
        output_content += "  Text:\n"
        for text in content['text']:
            output_content += f"    - {text}\n"
        if content['tables']:  # Only add tables section if there are tables
            output_content += "  Tables:\n"
            for table in content['tables']:
                output_content += "    Table:\n"
                for row in table:
                    output_content += f"      {row}\n"
        output_content += "\n" if content['text'] or content['tables'] else ""
    return output_content

def getLLamaResponse(prompt):
    
    client = Groq(
        api_key='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',http_client=httpx.Client(verify=False))
    
    completion = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {
                "role": "user",
                "content": prompt,
            }
        ],
        temperature=0.08,
        max_completion_tokens=32768,
        top_p=1,
        stream=False,
        stop=None,
    )

    return completion.choices[0].message.content



def getGeminiResponse(prompt):

    api_key = "AIzaSyB71ZnE68iS1fBjen5odVo_3BdkalMfihI"
    genai.configure(api_key=api_key,transport='rest')
    model = genai.GenerativeModel('models/gemini-2.0-flash')

    response = model.generate_content(contents=prompt)
    return response.text



def getResponse(content):
    
    API_KEY = "8MdaE49tRwKpn3HbkiwmQetoVrX8VRHP"
    url = "https://api.mistral.ai/v1/chat/completions"

    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }

    data = {
        "model": "mistral-tiny",
        "messages": [{"role": "user", "content": content}]
    }

    response = requests.post(url, headers=headers, json=data,verify=False)
    return response.json()["choices"][0]["message"]["content"]


@app.route('/upload', methods=['POST'])
def upload_file():
    file = request.files['file']
    if file:
        requirementsByHeading = extract_requirements_and_tables(file)
        requirements = extract_performance_requirements(file,requirementsByHeading)
        print(jsonify(requirements))
        return jsonify(requirements)
    return jsonify({"error": "No file uploaded"}), 400

if __name__ == '__main__':
    app.run(debug=True)
