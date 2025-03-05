import os
import PyPDF2 as pdf
import google.generativeai as genai
from docx import Document
 
 
#Step 1: LLM Connection setup
api_key = "AIzaSyB71ZnE68iS1fBjen5odVo_3BdkalMfihI"
genai.configure(api_key=api_key,transport='rest')
 
# specifying which model to use
model = genai.GenerativeModel('models/gemini-2.0-flash')
 
#Step 2: Extracts text from a URS file using python-docx.
def extract_from_docx(document_path):
   
    try:
        document = Document(document_path)
        full_text = []
        for paragraph in document.paragraphs:
            full_text.append(paragraph.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
 
        return None
   
 
def extract_Text(document_path):
 
    text_content = None
    # Plain text files
    if document_path.lower().endswith('.txt'):
        # print('txt document identified')
        text_content = extract_from_txt(document_path)
   
    # Word Documents
    elif document_path.lower().endswith('.docx'):
        # print('word document identified')
        text_content = extract_from_docx(document_path)
       
    # PDF files
    elif document_path.lower().endswith('.pdf'):
        # print('pdf identified')
        text_content = extract_from_pdf(document_path)
        #print(text_content)
   
    else:
        print("Unsupported document format. Please use .txt, .docx, or .pdf.")
        return None
   
    return text_content
 
def extract_Text_Upload(document_path):
 
    text_content = None
    # Plain text files
    if document_path.name.lower().endswith('.txt'):
        print('txt document identified')
        text_content = extract_from_txt(document_path)
   
    # Word Documents
    elif document_path.name.lower().endswith('.docx'):
        # print('word document identified')
        text_content =extract_from_docx(document_path)
       
    # PDF files
    elif document_path.name.lower().endswith('.pdf'):
        # print('pdf identified')
        text_content = extract_from_pdf(document_path)
        #print(text_content)
   
    else:
        print("Unsupported document format. Please use .txt, .docx, or .pdf.")
        return None
   
    return text_content
 
#Step 3: Extract content from local files for Few-shot prompting
 
URS_1 = extract_Text("C:/Users/pushpraj.sharma/Downloads/DocxExtraction/doc2/Sample_URS_001_PVCS.docx")
TS_1 = extract_Text("C:/Users/pushpraj.sharma/Downloads/DocxExtraction/doc2/Sample_TS_001_PVCS.docx")
URS_2 = extract_Text("C:/Users/pushpraj.sharma/Downloads/DocxExtraction/doc2/Sample_URS_004_SAVI.docx")
TS_2 = extract_Text("C:/Users/pushpraj.sharma/Downloads/DocxExtraction/doc2/Sample_TS_004_SAVI.docx")
 
def extract_from_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return text
    except Exception as e:
        print(f"Error reading text file: {e}")
        return None
 
# Extracting text from PDF
def extract_from_pdf(upload_file):
    reader = pdf.PdfReader(upload_file)
    text = ""
    for page in range(len(reader.pages)):
        page = reader.pages[page]
        text += str(page.extract_text())
    return text
 
# Step 4:  Extracting Perfromance Requirements and Generating Test case Using Gemini by Prompt Engineering
def analyze_text(doc_path):
 
    text_content =extract_Text(doc_path)
   
    prompt = f"""
   
    You are a QA Automation Engineer specialized in writing test cases for the User Acceptance testing.
 
    The {text_content} is extracted directly from the User Requirement Specification(URS) document for User Acceptance Testing(UAT).
 
    Analyse the following pairs of URS and Test Case texts:
 
    URS : {URS_1}
    Test Cases : {TS_1}
 
    URS : {URS_2}
    Test Cases : {TS_2}
 
    Using the above pair of URS and Test cases as reference for few shot learning, generate detailed test cases for {text_content} URS.
 
    Each test case should include:
        1. Test case ID
        2. Description of the test
        3. Step-by-step instructions
        4. Expected results
        5. Pass/fail criteria
        6. Traceability to the specific URS requirement
   
    Generate comprehensive UAT test cases.
           
    Please Format your responses clearly in a table.
       
    """
    try:
        response = model.generate_content(prompt)
        gemini_output = response.text
        # save_to_word(response.text,"gemini_output.docx")
       
        return gemini_output
 
    except Exception as e:
        print(f"Error during Gemini API call: {e}")
        return None
 
# Step 5: Write the gemini output to word file
def save_to_word(text, output_file):
 
    """
    Saves the Gemini response to a Word document.
    """
    try:
        # Initialize the Gemini model (configure your API key)
 
        # Create a Word document and add the generated text
        document = Document()
        document.add_paragraph(text)
        document.save(output_file)
 
        print(f"Gemini response saved to {output_file}")
 
    except Exception as e:
        print(f"An error occurred: {e}")

doc_path = "C:/Users/pushpraj.sharma/Downloads/DocxExtraction/Sample URS_004 SAVI.docx"
print("Here")
gemini_output = analyze_text(doc_path)
print(gemini_output)
 
# Step 6: Integrate the file upload and output into streamlit UI
"""
 
# Extracting from txt files
def extract_text_from_txt(file_path):
    Extracts text from a plain text (.txt) file.
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return text
    except Exception as e:
        print(f"Error reading text file: {e}")
        return None
 
        # Function to extract Performance Requirements/Qualifications based on simple keyword searches for headings
 
 
def extract_performance_requirement(doc_path):
 
    doc = Document(doc_path)
   
    sections = {}
 
    current_heading = None
 
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            current_heading = para.text  # Set current heading to heading text
            sections[current_heading] = []   # Initialize list for paragraphs under this heading
        elif current_heading:# If there is a current heading (meaning we are after a heading)
            sections[current_heading].append(para.text) # Add paragraph text to current heading's list
 
   
 
    output = []
 
    for heading, paragraphs in sections.items():
   
        if 'Performance Requirements' or 'PQ' or 'Performance Qualification' in heading:
            if paragraphs: # Check if there are paragraphs under this heading
                for p_text in paragraphs:
                    print(f"  - {p_text}")
                    output.append(p_text)
            else:
                print("  - (No paragraphs under this heading)")
 
 
    return '\n'.join(output)
 
def extract_text(document_path):
    print(document_path.name)
 
    text_content = None
    # Plain text files
    if document_path.name.lower().endswith('.txt'):
        print('txt document identified')
        text_content = extract_text_from_txt(document_path)
   
    # Word Documents
    elif document_path.name.lower().endswith('.docx'):
        print('word document identified')
        text_content = extract_text_from_docx(document_path)
       
    # PDF files
    elif document_path.name.lower().endswith('.pdf'):
        print('pdf identified')
        text_content = extract_text_from_pdf(document_path)
        print(text_content)
   
    else:
        print("Unsupported document format. Please use .txt, .docx, or .pdf.")
        return None
 
 
"""